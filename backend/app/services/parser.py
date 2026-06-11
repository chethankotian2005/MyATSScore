import io
from fastapi import UploadFile, HTTPException, status
import fitz  # PyMuPDF
import docx

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

class ResumeParser:
    def __init__(self, file: UploadFile):
        self.file = file
        self.filename = file.filename.lower() if file.filename else ""
        self.raw_text = ""
        self.sections = {
            "contact": [],
            "summary": [],
            "experience": [],
            "education": [],
            "skills": [],
            "certifications": [],
            "projects": []
        }
        self.has_tables = False
        self.has_columns = False
        self.has_images = False

    async def parse(self) -> dict:
        content = await self.file.read()
        
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="File exceeds 5MB limit."
            )
            
        if self.filename.endswith('.pdf'):
            self._parse_pdf(content)
        elif self.filename.endswith('.docx'):
            self._parse_docx(content)
        elif self.filename.endswith('.txt'):
            self._parse_txt(content)
        else:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="Unsupported file format. Please upload PDF or DOCX."
            )
            
        self._extract_sections()
        
        return {
            "raw_text": self.raw_text.strip(),
            "sections": self.sections,
            "word_count": len(self.raw_text.split()),
            "has_tables": self.has_tables,
            "has_columns": self.has_columns,
            "has_images": self.has_images
        }

    def _parse_pdf(self, content: bytes):
        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Corrupted file."
            )
            
        if doc.needs_pass:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Password-protected PDFs are not supported."
            )
            
        for page in doc:
            if page.get_image_info():
                self.has_images = True
                
            try:
                if page.find_tables().tables:
                    self.has_tables = True
            except Exception:
                pass
                
            blocks = page.get_text("blocks")
            
            x0_positions = set()
            for b in blocks:
                if len(b) >= 7 and b[6] == 0:  # text block
                    x0_positions.add(round(b[0], -1))
                elif len(b) >= 7 and b[6] == 1: # image block
                    self.has_images = True
            
            if len(x0_positions) > 3: # multiple margins = columns
                self.has_columns = True
            
            self.raw_text += page.get_text("text") + "\n"

    def _parse_docx(self, content: bytes):
        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Corrupted file."
            )
            
        if len(doc.tables) > 0:
            self.has_tables = True
            
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                self.has_images = True
                break
                
        self.has_columns = False 
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        self.raw_text = "\n".join(full_text)
        
    def _parse_txt(self, content: bytes):
        try:
            self.raw_text = content.decode('utf-8')
        except UnicodeDecodeError:
            self.raw_text = content.decode('latin-1', errors='ignore')
            
        self.has_tables = False
        self.has_columns = False
        self.has_images = False
        
    def _extract_sections(self):
        lines = self.raw_text.split('\n')
        current_section = "contact"
        
        section_keywords = {
            "summary": ["summary", "profile", "objective", "about me", "professional summary"],
            "experience": ["experience", "employment", "work history", "professional experience", "work experience"],
            "education": ["education", "academic background", "degrees", "academic history"],
            "skills": ["skills", "technologies", "core competencies", "technical skills", "expertise"],
            "certifications": ["certifications", "licenses", "certificates"],
            "projects": ["projects", "personal projects", "academic projects", "open source"]
        }
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
                
            line_lower = line_clean.lower()
            
            is_new_section = False
            words = line_lower.split()
            if len(words) <= 4:
                clean_header = line_lower.replace(":", "").strip()
                for sec, keywords in section_keywords.items():
                    if clean_header in keywords:
                        current_section = sec
                        is_new_section = True
                        break
                        
            if not is_new_section:
                self.sections[current_section].append(line_clean)
