import pytest
import io
from fastapi import HTTPException
from app.services.parser import ResumeParser
import fitz
import docx

class DummyUploadFile:
    def __init__(self, filename: str, content: bytes):
        self.filename = filename
        self._content = content
        
    async def read(self) -> bytes:
        return self._content

@pytest.mark.asyncio
async def test_parse_unsupported_format():
    file = DummyUploadFile("resume.txt", b"some text")
    parser = ResumeParser(file)
    with pytest.raises(HTTPException) as exc_info:
        await parser.parse()
    assert exc_info.value.status_code == 415

@pytest.mark.asyncio
async def test_parse_file_too_large():
    # 5MB + 1 byte
    large_content = b"0" * (5 * 1024 * 1024 + 1)
    file = DummyUploadFile("resume.pdf", large_content)
    parser = ResumeParser(file)
    with pytest.raises(HTTPException) as exc_info:
        await parser.parse()
    assert exc_info.value.status_code == 413
    assert "5MB limit" in exc_info.value.detail

@pytest.mark.asyncio
async def test_parse_corrupted_pdf():
    file = DummyUploadFile("corrupted.pdf", b"not a real pdf content here")
    parser = ResumeParser(file)
    with pytest.raises(HTTPException) as exc_info:
        await parser.parse()
    assert exc_info.value.status_code == 400
    assert "Corrupted file" in exc_info.value.detail

@pytest.mark.asyncio
async def test_parse_docx_success():
    doc = docx.Document()
    doc.add_paragraph("Summary")
    doc.add_paragraph("A great software engineer.")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Worked at Google.")
    
    stream = io.BytesIO()
    doc.save(stream)
    content = stream.getvalue()
    
    file = DummyUploadFile("resume.docx", content)
    parser = ResumeParser(file)
    result = await parser.parse()
    
    assert "A great software engineer." in result["sections"]["summary"]
    assert "Worked at Google." in result["sections"]["experience"]
    assert result["has_tables"] is False
    assert result["has_columns"] is False
    assert result["has_images"] is False

@pytest.mark.asyncio
async def test_parse_pdf_success():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Education\nB.S. Computer Science\nSkills\nPython, FastAPI")
    content = doc.write()
    doc.close()
    
    file = DummyUploadFile("resume.pdf", content)
    parser = ResumeParser(file)
    result = await parser.parse()
    
    assert "B.S. Computer Science" in result["sections"]["education"]
    assert "Python, FastAPI" in result["sections"]["skills"]
    assert result["has_tables"] is False
